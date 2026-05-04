"""
File Parser — Extracts raw text from PDF and DOCX files.

Supports:
  - PDF via PyMuPDF (fitz)
  - DOCX via python-docx
"""

import os
from pathlib import Path
from langchain_core.documents import Document


def parse_pdf(file_path: str) -> list[Document]:
    """Extract text from a PDF file, one Document per page."""
    import fitz  # PyMuPDF

    documents = []
    pdf = fitz.open(file_path)
    filename = os.path.basename(file_path)

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text")
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": filename,
                        "page": page_num + 1,
                        "total_pages": len(pdf),
                    },
                )
            )
    pdf.close()
    return documents


def parse_docx(file_path: str) -> list[Document]:
    """Extract text from a DOCX file, one Document per logical section."""
    from docx import Document as DocxDocument

    docx = DocxDocument(file_path)
    filename = os.path.basename(file_path)
    full_text = []
    current_section = ""
    section_num = 0
    documents = []

    for para in docx.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headings (bold or heading styles)
        is_heading = (
            para.style.name.startswith("Heading")
            or (para.runs and para.runs[0].bold)
        )

        if is_heading and current_section:
            # Save previous section
            section_num += 1
            documents.append(
                Document(
                    page_content=current_section,
                    metadata={
                        "source": filename,
                        "section": section_num,
                        "heading": full_text[-1] if full_text else "",
                    },
                )
            )
            current_section = text + "\n"
            full_text = [text]
        else:
            current_section += text + "\n"
            full_text.append(text)

    # Save last section
    if current_section.strip():
        section_num += 1
        documents.append(
            Document(
                page_content=current_section,
                metadata={
                    "source": filename,
                    "section": section_num,
                },
            )
        )

    # Fallback: if no sections were detected, return as a single document
    if not documents and full_text:
        documents.append(
            Document(
                page_content="\n".join(full_text),
                metadata={"source": filename, "section": 1},
            )
        )

    return documents


def parse_file(file_path: str) -> list[Document]:
    """
    Route to the appropriate parser based on file extension.

    Args:
        file_path: Path to the uploaded file (PDF or DOCX).

    Returns:
        List of LangChain Document objects with metadata.

    Raises:
        ValueError: If the file type is unsupported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return parse_docx(file_path)
    elif ext == ".txt":
        # Simple fallback for plain text
        text = path.read_text(encoding="utf-8")
        return [
            Document(
                page_content=text,
                metadata={"source": path.name, "page": 1},
            )
        ]
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. Supported formats: .pdf, .docx, .txt"
        )
